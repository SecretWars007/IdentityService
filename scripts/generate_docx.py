import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
output_dir = os.path.join(root, 'docs')
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'IdentityService_Documentation.docx')

exclude_dirs = {'bin', 'obj', '.git', '.vs'}

cs_files = []
for dirpath, dirnames, filenames in os.walk(root):
    # skip excluded dirs
    parts = set(dirpath.split(os.sep))
    if parts & exclude_dirs:
        continue
    for f in filenames:
        if f.endswith('.cs'):
            cs_files.append(os.path.join(dirpath, f))

cs_files.sort()

doc = Document()
doc.add_heading('IdentityService - Documentación del código', level=1)

# add a code style
styles = doc.styles
if 'Code' not in styles:
    code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    font = code_style.font
    font.name = 'Consolas'
    font.size = Pt(9)
else:
    code_style = styles['Code']

for path in cs_files:
    rel = os.path.relpath(path, root)
    doc.add_heading(rel.replace('\\', '/'), level=2)
    try:
        with open(path, 'r', encoding='utf-8') as fh:
            text = fh.read()
    except Exception:
        with open(path, 'r', encoding='latin-1') as fh:
            text = fh.read()

    # namespace
    ns_match = re.search(r'namespace\s+([\w\.]+)', text)
    namespace = ns_match.group(1) if ns_match else 'No se detectó namespace'
    doc.add_paragraph(f'Namespace: {namespace}')

    # try to find top-level summary comment (///) or /* */ at top
    summary = None
    xml_comments = re.findall(r'(///.*(?:\n///.*){0,9})', text)
    if xml_comments:
        # take first block and clean
        block = xml_comments[0]
        lines = [l.strip()[3:].strip() for l in block.splitlines()]
        summary = ' '.join(lines).strip()
    else:
        m = re.search(r'/\*\*(.*?)\*/', text, re.S)
        if m:
            summary = ' '.join(l.strip('* ').strip() for l in m.group(1).splitlines()).strip()

    if summary:
        doc.add_paragraph('Resumen (autogenerado): ' + summary)
    else:
        doc.add_paragraph('Resumen (autogenerado): No disponible')

    # detect classes, interfaces, enums
    classes = re.findall(r'\bclass\s+(\w+)', text)
    interfaces = re.findall(r'\binterface\s+(\w+)', text)
    enums = re.findall(r'\benum\s+(\w+)', text)

    if classes:
        doc.add_paragraph('Clases: ' + ', '.join(classes))
    if interfaces:
        doc.add_paragraph('Interfaces: ' + ', '.join(interfaces))
    if enums:
        doc.add_paragraph('Enums: ' + ', '.join(enums))

    doc.add_paragraph('Código fuente:', style=None)
    # add code block preserving lines
    lines = text.splitlines()
    for ln in lines:
        p = doc.add_paragraph(ln, style='Code')

    doc.add_page_break()

# add a footer page with stats
doc.add_heading('Resumen del repositorio', level=1)
doc.add_paragraph(f'Archivos .cs documentados: {len(cs_files)}')

doc.save(output_path)
print(f'Document generado: {output_path}')
